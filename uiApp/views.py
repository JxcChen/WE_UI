import json
import platform
import re
import shutil
import subprocess
import threading
import time
import zipfile

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, HttpResponseRedirect
from django.shortcuts import render
from uiApp.models import *
from docx import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.contrib import auth
from django.contrib.auth.models import User

# Create your views here.

"""
视图逻辑层
"""
operation = platform.system()


def to_login(request):
    next_url = request.GET.get('next')
    return render(request, 'login.html', {'next': next_url})


def login(request):
    next_url = request.GET.get('next', '/home/')
    print(next_url)
    if next_url == '' or next_url == 'None':
        next_url = '/home/'
    username = request.POST.get('username')
    password = request.POST.get('password')

    # authenticate 有就返回对象 没有就返回none
    username_check = len(User.objects.filter(username=username))
    if username_check == 0:
        return HttpResponseRedirect("/login_no_user/")
    user = auth.authenticate(username=username, password=password)
    if not user:
        return HttpResponseRedirect("/login_pwd_error/")
    auth.login(request, user)
    request.session['user'] = username
    return HttpResponseRedirect(next_url)


def login_no_user(request):
    return render(request, 'login.html', {'error_msg': "该用户不存在"})


def login_pwd_error(request):
    return render(request, 'login.html', {'error_msg': "密码错误"})


# 退出功能
def logout(request):
    auth.logout(request)
    return HttpResponseRedirect('/login/')


def user_list(request):
    users = list(User.objects.all().values())
    return render(request, 'user.html', {'users': users})


def add_user(request):
    username = request.POST.get('username')
    pwd = request.POST.get('password')
    try:
        user = User.objects.create_user(username=username, password=pwd)
        user.save()
    except:
        return HttpResponse('用户名已存在')
    return HttpResponseRedirect('/user_list/')


# 需要先修改setting.py 才有home.html联想
@login_required()
def home(request):
    res = {'end': DB_end.objects.all(), 'href': DB_href.objects.all()}
    return render(request, 'home.html', res)


# 保存项目
def save_end(request):
    pro_name = request.GET['name']
    pro_host = request.GET['host']
    pro_email = request.GET['email']
    end = DB_end.objects.create(name=pro_name, host=pro_host, email=pro_email)
    base_path = os.path.dirname(os.path.abspath(__file__))
    demo_path = os.path.join(base_path, r"../my_client/demo_client")
    new_client = os.path.join(base_path, r"my_client/client_" + str(end.id))
    shutil.copytree(demo_path, new_client)
    return HttpResponse("")


# 删除项目
def del_end(request, del_id):
    DB_end.objects.filter(id=del_id).delete()
    return HttpResponseRedirect("/home/")


# 获取项目信息
def get_project_msg(request, pro_id):
    pro_msg = list(DB_end.objects.filter(id=pro_id).values())[0]
    return HttpResponse(json.dumps(pro_msg), content_type="application/json")


# 编辑项目
def update_project(request):
    DB_end.objects.filter(id=request.GET['pro_id']).update(name=request.GET['pro_name'], host=request.GET['pro_host'],
                                                           monitor_host=request.GET['monitor_host'],
                                                           check_time=request.GET['check_time'],
                                                           phone=request.GET['phone'], email=request.GET['email'],
                                                           dingtalk=request.GET['dingtalk'],
                                                           max_threads=request.GET['max_threads'], )
    return HttpResponseRedirect("/home/")


# 进入项目测试用例页面
@login_required
def testcases(request, pro_id):
    case_name = request.GET.get('case_name', '')  # 获取搜索的用例名称
    cases = list(DB_cases.objects.filter(pro_id=pro_id, name__contains=case_name))
    # 将cases进行分页处理  每页5条数据
    p = Paginator(cases, 5)
    # 获取前端传回来的页数
    current_page = request.GET.get('current_page')
    try:
        page_cases = p.page(current_page)  # 获取当前页的测试用例
    except PageNotAnInteger:
        page_cases = p.page(1)  # 页数非整数时直接返回第一页
    except EmptyPage:
        page_cases = p.page(1)  # 页数为空时返回第一页
    project = list(DB_end.objects.filter(id=pro_id))[0]
    hosts = project.host.split(",")
    param = {"cases": page_cases, "project": project, "hosts": hosts, 'key_word': case_name}
    return render(request, 'case.html', param)


# 添加测试用例
def add_case(request, pro_id):
    is_monitor = False
    if request.GET['is_monitor'] == "True":
        is_monitor = True

    is_threads = False
    if request.GET['is_monitor'] == "True":
        is_monitor = True

    DB_cases.objects.create(pro_id=pro_id, name=request.GET['case_name'], retry_count=request.GET['retry_count'],
                            is_monitor=is_monitor, is_threads=is_threads, case_type=request.GET['case_type'])

    return HttpResponseRedirect('/testcases/' + pro_id + '/')


# 编辑测试用例
def edit_case(request):
    case_id = request.GET['id']
    case = list(DB_cases.objects.filter(id=case_id).values())[0]
    return HttpResponse(json.dumps(case), content_type="application/json")


# 更新测试用例
def update_case(request, pro_id):
    is_monitor = False
    if request.GET['is_monitor'] == "True":
        is_monitor = True

    is_threads = False
    if request.GET['is_threads'] == "True":
        is_threads = True
    case_id = request.GET['case_id']

    DB_cases.objects.filter(id=case_id).update(name=request.GET['case_name'], retry_count=request.GET['retry_count'],
                                               is_monitor=is_monitor, is_threads=is_threads,
                                               case_type=request.GET['case_type'])
    return HttpResponse('')


# 上传脚本
def upload_script(request, case_id):
    # 拿到端id
    pro_id = DB_cases.objects.filter(id=case_id)[0].pro_id
    # 获取到上传的脚本
    my_file = request.FILES.get("script_file", None)
    # 如果文件为空则直接返回
    if not my_file:
        return HttpResponseRedirect('/testcases/' + pro_id + '/')
    # 获取文件名称
    file_name = str(my_file)
    # 打开本地文件
    with open("my_client/client_%s/case/%s" % (pro_id, file_name), 'wb') as file:
        for content in my_file.chunks():
            file.write(content)
    # 将文件名称存到数据库
    DB_cases.objects.filter(id=case_id).update(script=file_name)
    # 返回
    return HttpResponseRedirect('/testcases/%s/' % pro_id)


# 删除用例
def del_case(request):
    case_id = request.GET['case_id']
    pro_id = request.GET['pro_id']
    DB_cases.objects.filter(id=case_id).delete()

    return HttpResponseRedirect('/testcases/' + pro_id + '/')


# 运行脚本
def run_script(request, case_id):
    case = DB_cases.objects.filter(id=case_id)[0]
    pro_id = case.pro_id
    case_name = case.name
    script_name = case.script
    host = request.GET['host']

    # 判断是否未上传脚本文件
    if script_name in ['', ' ', None, 'None']:
        return HttpResponse('Error')
    # 执行py文件
    if operation == "Windows":
        subprocess.call(
            'python my_client/client_%s/case/%s %s %s %s' % (pro_id, script_name, host, script_name, case_name),
            shell=True)
    else:
        subprocess.call(
            'python3 my_client/client_%s/case/%s %s %s %s' % (pro_id, script_name, host, script_name, case_name),
            shell=True)
    return HttpResponse('Success')


# 并发执行脚本
def concurrent_run_script(request, pro_id):
    # 先获取该项目下所有可以并发执行的测试用例
    cases = DB_cases.objects.filter(pro_id=pro_id, is_threads='True')
    max_threads = DB_end.objects.filter(id=pro_id)[0].max_threads
    host = request.GET['host']

    # 声明执行用例方法
    def concurrent_run(case):
        if case.script not in ['', ' ', None, 'None']:
            # 执行py文件
            if operation == "Windows":
                subprocess.call('python my_client/client_%s/case/%s %s %s %s' % (
                    case.pro_id, case.script, host, case.script, case.name), shell=True)
            else:
                subprocess.call('python3 my_client/client_%s/case/%s %s %s %s' % (
                    case.pro_id, case.script, host, case.script, case.name), shell=True)
            print(case, "执行完成")

    tf = []

    # 声明多个线程执行运行用例方法
    for case in cases:
        t = threading.Thread(target=concurrent_run, args=(case,))
        t.setDaemon(True)  # 将线程声明为守护线程
        tf.append(t)
    # 限制最大并发

    for i in range(0, len(tf), max_threads):
        tmp = tf[i:i + max_threads]
        # 执行
        for t in tmp:
            t.start()  # 运行线程任务

        for t in tmp:
            t.join()  # 子线程再未完成的情况下 主线程会一直处于阻塞状态
    return HttpResponse('Success')


# 开启监控
def open_monitor(request, pro_id):
    monitor_host = DB_end.objects.filter(id=pro_id)[0].monitor_host
    # 判断监控是否已经开启
    try:
        process = subprocess.check_output('ps -ef | grep "monitor.py %s WEB" | grep -v grep' % pro_id, shell=True)
        print("监控已经开启")
    except:
        # 未开启则开启监控
        def start_monitor():
            if operation == 'Windows':
                subprocess.call('python uiApp/monitor.py %s WEB' % pro_id, shell=True)
            else:
                subprocess.call('python3 uiApp/monitor.py %s WEB' % pro_id, shell=True)

        # 开一个新的线程进行监控
        t = threading.Thread(target=start_monitor)
        t.setDaemon(True)  # 设置守护进程
        # 执行进程
        t.start()
    return HttpResponseRedirect("/testcases/%s/" % pro_id)


# 关闭监控
def close_monitor(request, pro_id):
    # 查看监控是否开启 开启则关闭
    try:
        process = subprocess.check_output('ps -ef | grep "monitor.py %s WEB" | grep -v grep' % pro_id, shell=True)
        print("监控已开启，现在关闭")
        # 使用正则
        pids = re.findall(r'(\d+)', str(process))
        pid = max([int(i) for i in pids])
        subprocess.call('kill -9 %s' % pid,
                        shell=True)
    except:
        print("监控尚未开启")
    return HttpResponseRedirect("/testcases/%s/" % pro_id)


# 查看报告
def look_report(request, case_id):
    case = DB_cases.objects.filter(id=int(case_id))[0]
    case_name = case.name
    # 返回测试报告路径
    return render(request, 'client_%s/report/%s.html' % (case.pro_id, case_name))


# 查看报告总结
@login_required
def look_report_summary(request, pro_id=''):
    # 获取项目对应的全部用例
    pro_name = DB_end.objects.filter(id=pro_id)[0].name
    cases = list(DB_cases.objects.filter(pro_id=pro_id).values())
    # 声明结果变量
    res = '【%s项目用例总结】\n' % pro_name
    total_case = 0;
    pass_case = 0
    fail_case = 0
    # 存放错误用例名称
    fail_case_list = []
    # 遍历用例报告获取总结数据
    for case in cases:
        try:
            with open(r'my_client/client_%s/report/%s.html' % (pro_id, case['name']), 'r', encoding='utf-8') as f:
                # 读取报告内容
                content = f.read()
                # 使用正则匹配结果
                results = re.findall(r"<td name='sum'>(.*?)</td>", content)
                total_case += int(results[0])
                pass_case += int(results[1])
                fail_or_error = int(results[2]) + int(results[3])
                fail_case += fail_or_error
                if fail_or_error > 0:
                    fail_case_list.append(case['name'])
        except FileNotFoundError as e:
            # 如果没找到文件的话表示该用例还未执行  直接跳过即可
            continue
        except Exception as e:
            raise e
    res += "当前总共有【%s】条用例。\n通过用例数：%s 失败用例数：%s\n" % (str(total_case), str(pass_case), str(fail_case))
    res += "失败用例名称：" + ",".join(fail_case_list) + "\n"
    res += "想查看用例结果详情可以点击用例后的报告按钮"
    return HttpResponse(res)


# 导出测试报告 word格式
def export_report(request, pro_id):
    # 直接调用获取报告函数
    httpresponse = look_report_summary(request, pro_id=pro_id)
    project = DB_end.objects.filter(id=pro_id)[0]

    # 创建word文档
    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'

    # 写标题
    p = doc.add_paragraph('【%s 的自动化测试报告】' % project.name)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

    p = doc.add_paragraph('')
    p.add_run('本测试报告由xxx的web自动化平台:xxxxxx生成,地址:xxxxxx').font.color.rgb = RGBColor(*(172, 182, 182))
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

    p = doc.add_paragraph('')
    p.add_run('报告生成时间: %s' % time.strftime('%Y-%m-%d %H:%M:%s')).font.color.rgb = RGBColor(*(172, 182, 182))
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

    doc.add_paragraph('报告人: %s' % request.user.username, style='Intense Quote')
    doc.add_paragraph('平台维护: %s' % '测试开发干货', style='Intense Quote')
    doc.add_paragraph('执行环境: %s' % 'http://', style='Intense Quote')

    doc.add_paragraph('用例执行结果总览如下:', style='Intense Quote')

    p = doc.add_paragraph(httpresponse.content.decode())
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('用例执行结果传送门如下:', style='Intense Quote')

    # 获取所有报告的详情地址并展示到word中
    cases = DB_cases.objects.filter(pro_id=pro_id)
    for i in cases:
        doc.add_paragraph('用例名字:' + str(i.name))
        doc.add_paragraph('脚本名字:' + str(i.script))
        doc.add_paragraph('报告地址:' + 'http://xxxx:8000/look_report/' + str(i.id) + '/')
        doc.add_paragraph(' ', style='Intense Quote')

    # 保存并下载
    file_name = 'TESTREPORTS_%s.docx' % time.strftime('%Y-%m-%d %H:%M:%s')
    doc.save(file_name)
    with open(file_name, 'rb') as fp:
        response = HttpResponse(fp)
        response['Content-Type'] = 'application/octet-stream'
        response['Content-Disposition'] = 'attachment;filename="%s"' % file_name

    os.remove(file_name)

    return response


# 压缩函数
def get_zip_file(input_path, result):
    """
    对目录进行深度优先遍历
    :param input_path:
    :param result:
    :return:
    """
    files = os.listdir(input_path)
    for file in files:
        if os.path.isdir(input_path + '/' + file):
            get_zip_file(input_path + '/' + file, result)
        else:
            result.append(input_path + '/' + file)


def zip_file_path(input_path, output_path, output_name):
    """
    压缩文件
    :param input_path: 压缩的文件夹路径
    :param output_path: 解压（输出）的路径
    :param output_name: 压缩包名称
    :return:
    """
    f = zipfile.ZipFile(output_path + '/' + output_name, 'w', zipfile.ZIP_DEFLATED)
    filelists = []
    get_zip_file(input_path, filelists)
    for file in filelists:
        f.write(file)
    # 调用了close方法才会保证完成压缩
    f.close()


# 下载本地调试包
def download_client(request, pro_id):
    # 1 获取到本地调试包名称
    zip_file_name = "client_%s.zip" % pro_id
    file_dir = "my_client/client_%s" % pro_id
    # 2 删除旧的调试压缩包
    try:
        os.remove("my_client/" + zip_file_name)
    except:
        pass
    # 3 将最新的调试包进行打包压缩  调用函数
    zip_file_path(file_dir, "my_client", zip_file_name)
    # 4 封装响应  响应体 头
    # 先以读方式打开文件
    try:
        file = open("my_client/" + zip_file_name, 'rb')
    except:
        # 没有找到代表没有这个文件
        return HttpResponse('')
    response = HttpResponse(file)
    response['Content-Type'] = 'application/octet-stream'
    response['Content-Disposition'] = 'attachment;filename="%s"' % zip_file_name
    # 5 删除调试压缩包
    try:
        os.remove("my_client/" + zip_file_name)
    except:
        pass
    # 6 返回响应
    return response


# 上传public-utils
def upload_public_utils(request, pro_id):
    # 获取页面传回来的文件
    utils_file = request.FILES.get("utils_file", None)
    file_name = str(utils_file)
    if not utils_file:
        return HttpResponseRedirect('/testcases/' + pro_id)
    new_file_name = 'my_client/client_%s/public/%s' % (pro_id, file_name)
    # 读取传回来的文件写到新文件中
    with open(new_file_name, 'wb+') as f:
        for content in utils_file.chunks():
            f.write(content)
    return HttpResponseRedirect('/testcases/' + pro_id)
